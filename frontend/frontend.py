import streamlit as st
import plotly.graph_objects as go
from keras.models import load_model
from tkinter import Tk
from tkinter.filedialog import askopenfilename

import joblib
from joblib import load
import numpy as np
import streamlit as st
import PyPDF2
from PyPDF2 import PdfReader
from docx import Document
import time


# Laden der Imports für die Speech to Text Funktion
import speech_recognition as sr
from time import ctime # get time details
import time
import playsound
import os
import random
from gtts import gTTS

from python_scripts.transcript_creation import modell_speak, record_audio

##############################################
# Laden der Imports für die Zusammenfassung
import re
import networkx as nx
import spacy
from summa import keywords
from summa.summarizer import summarize

# Laden des Spacy-Modells
import evaluate
import nltk
from nltk.tokenize import sent_tokenize
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer,PreTrainedTokenizerFast
import matplotlib.pyplot as plt
import math
from tqdm import tqdm
from transformers import pipeline
import torch
from deepmultilingualpunctuation import PunctuationModel

from python_scripts.zusammenfassung import execute_text_gen
##############################################


############ Laden der Modelle ###############
model_text_korrigieren = PunctuationModel()
## Laden des Modells für die Klassifikation
model_classification = load_model("classification_modells/neuro_net_1.h5")
vectorizer = load("classification_modells/vectorizer_1.joblib")

## Laden des Modells für die Zusammenfassung
for model_name in ['NICFRU/bart-base-paraphrasing-science','NICFRU/bart-base-paraphrasing-news','NICFRU/bart-base-paraphrasing-story','NICFRU/bart-base-paraphrasing-review']:
    tokenizer_zusammenfassung = AutoTokenizer.from_pretrained(model_name)
    paraphrase_zusammenfassung = pipeline("text2text-generation", model=model_name)


##############################################

# Schriftgröße für Buttons anpassen
st.write('<style>div.row-widget button{font-size: 30px !important;}</style>', unsafe_allow_html=True)

# Schriftgröße für Text Area anpassen
st.write('<style>div.row-widget textarea{font-size: 35px !important;}</style>', unsafe_allow_html=True)

m = st.markdown("""
<style>
div.stButton > button:nth-of-type(2) {
    background-color: #ce1126;
    color: white;
    height: 3em;
    width: 12em;
    border-radius:10px;
    border:3px solid #000000;
    font-size:20px;
    font-weight: bold;
    margin: auto;
    display: block;
}

div.stButton > button:hover {
	background:linear-gradient(to bottom, #ce1126 5%, #ff5a5a 100%);
	background-color:#ce1126;
}
    /* 1st button */
    .element-container:nth-child(3) {
      left: 10px;
      top: -60px;
    }
div.stButton > button:active {
	position:relative;
	top:3px;
}

</style>""", unsafe_allow_html=True)
# engine = pyttsx3.init()
# engine.setProperty('rate', 150)  # Anpassen der Sprechgeschwindigkeit (optional)

# def text_to_speech(text):
#     global engine
#     engine.say(text)
#     engine.runAndWait()  

def extract_text_from_element(element):
    # Extrahiere den Text aus einem Streamlit-Element
    if hasattr(element, 'text'):
        return element.text
    elif isinstance(element, (list, tuple)):
        return ' '.join([extract_text_from_element(item) for item in element])
    elif isinstance(element, dict):
        return ' '.join([extract_text_from_element(item) for item in element.values()])
    else:
        return str(element)
    
def extract_text_from_page(page):
    # Extrahiere den Text aus einer Streamlit-Seite
    text = ''
    for element in page:
        text += extract_text_from_element(element)
    return text

# # Texteingabe


# Button zum Auslösen der Sprachausgabe

def read_docx_file(file_path):
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return " ".join(paragraphs)

def read_txt_file(file_path):
    with open(file_path, 'r') as file:
        content = file.read()
    return content

def read_pdf_file(file_path):
    content = ""
    with open(file_path, 'rb') as file:
        pdf = PyPDF2.PdfReader(file)
        for page in pdf.pages:
            content += page.extract_text()
    return content

def read_pdf(file):
    data_folder = "Testfiles"
    file_path = os.path.join(data_folder, file)
    pdf = PdfReader(file_path)
    text = ""
    for page in range(len(pdf.pages)):
        text += pdf.pages[page].extract_text()
    return text

def read_docx(file):
    data_folder = "Testfiles"
    file_path = os.path.join(data_folder, file)
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_txt(file):
    data_folder = "Testfiles"
    file_path = os.path.join(data_folder, file.name)
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return text

def summarize_text(text, compression_rate,predicted_class):
    # Hier erfolgt die Zusammenfassung des Textes unter Berücksichtigung der Kompressionsrate
    # Implementiere hier deine Logik zur Zusammenfassung des Textes
    dic=execute_text_gen({'classification':predicted_class,'text':text,'compression':compression_rate*0.01})
    paraphrase_zusammenfassung = dic['text_rank_text_2']
    text_compression_rate=round(dic['ent_com_rate']*100,2)
    print('Text: ',paraphrase_zusammenfassung,'Kompression: ' ,text_compression_rate)
    return paraphrase_zusammenfassung, text_compression_rate

    #return summarized_text

def classify_text(text):
    new_text_features = vectorizer.transform([text])
    predictions = model_classification.predict(new_text_features.toarray())
    predicted_class = np.argmax(predictions, axis=1)
    predicted_probability = np.max(predictions, axis=1)
    return predicted_class, predicted_probability
def style_button_row(clicked_button_ix, n_buttons):
    def get_button_indices(button_ix):
        return {
            'nth_child': button_ix,
            'nth_last_child': n_buttons - button_ix + 1
        }

    clicked_style = """
    div[data-testid*="stHorizontalBlock"] > div:nth-child(%(nth_child)s):nth-last-child(%(nth_last_child)s) button {
        border-color: rgb(255, 75, 75);
        color: rgb(255, 75, 75);
        box-shadow: rgba(255, 75, 75, 0.5) 0px 0px 0px 0.2rem;
        outline: currentcolor none medium;
    }
    """
    unclicked_style = """
    div[data-testid*="stHorizontalBlock"] > div:nth-child(%(nth_child)s):nth-last-child(%(nth_last_child)s) button {
        pointer-events: none;
        cursor: not-allowed;
        opacity: 0.65;
        filter: alpha(opacity=65);
        -webkit-box-shadow: none;
        box-shadow: none;
    }
    """
    style = ""
    for ix in range(n_buttons):
        ix += 1
        if ix == clicked_button_ix:
            style += clicked_style % get_button_indices(ix)
        else:
            style += unclicked_style % get_button_indices(ix)
    st.markdown(f"<style>{style}</style>", unsafe_allow_html=True)
# Streamlit-App

def execute_the_classification_and_summary(input_text,compression_rate,classification, summary):
    predicted_class, predicted_probability = classify_text(input_text)
    predicted_probability = round(predicted_probability[0]*100, 2)
    if predicted_class == 0:
        predicted_class = "Scientific Paper"
    elif predicted_class == 1:
        predicted_class = "News"
    elif predicted_class == 2:
        predicted_class = "Review"
    else:
        predicted_class = "Story"
    

    if predicted_probability >= 95:
        if classification:
            st.subheader("Classification")
            st.write(f"The class calculated by the model is:  \"{predicted_class}\". The model is very confident with a probability of {predicted_probability}%.")
        if summary:
            paraphrase_zusammenfassung, text_compression_rate = summarize_text(input_text, compression_rate,predicted_class)
            st.subheader("Summary")
            st.write("This is a summary of the input text has a compression rate of {}%:".format(text_compression_rate))
            st.write(paraphrase_zusammenfassung)
    else:
        st.write(f"Attention! The model is not entirely certain. The class calculated by the model is: \"{predicted_class}\". However, the model predicts this class with a probability of only {predicted_probability}%.")

def split_into_batches(text, batch_size=500):
    # First, tokenize the text
    nltk.download('punkt')
    tokens = nltk.word_tokenize(text)

    # Split into batches
    batches = [tokens[i:i + batch_size] for i in range(0, len(tokens), batch_size)]
    
    # Join tokens back into strings
    batches = [' '.join(batch) for batch in batches]

    return batches
def is_file(path):
    return os.path.isfile(path)
def main():
    
    # Führe eine Wartezeit von wait_time Sekunden durch
    #time.sleep(30)
    global content
    try:
        if content != "":
            content=content
        else:
            content = ""
    except:
        content = ""
    # Texteingabe
    
    # text_vorlesen = "Welcome to Syntex, to proceed further with the screenreader function, tap the big red botton on the left in the middle of your screen"
    # modell_speak(text_vorlesen)


    st.title("SynTex")
    file = st.file_uploader("Please select a document", type=["pdf", "docx", "txt"])
    print(file)

    # Add custom CSS style
    st.markdown(
        """
        <style>
        .red-button {
            background-color: red;
            color: white;
            padding: 0.5rem 2rem;
            font-size: 1.5rem;
            border-radius: 0.5rem;
            border: none;
            cursor: pointer;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # if st.button("Upload a document",key=1,type='primary'):
    #     root = Tk()
    #     root.withdraw()
    #     file_path = askopenfilename(filetypes=[("Textdokumente", "*.txt"), ("Word-Dokumente", "*.docx"), ("PDF-Dateien", "*.pdf")])

    #     if file_path:
    #         if file_path.endswith('.docx'):
    #             input_text = read_docx_file(file_path)
    #             print(input_text)
    #         elif file_path.endswith('.txt'):
    #             input_text = read_txt_file(file_path)
    #             print(input_text)
    #         elif file_path.endswith('.pdf'):
    #             input_text = read_pdf_file(file_path)
    #             print(input_text)
    #         else:
    #             print("Ungültiger Dateityp. Nur DOCX-, TXT- und PDF-Dateien werden unterstützt.")
    #     else:
    #         print("Keine Datei ausgewählt.")
    
    


    if file is not None:
        
        content = ""
        
        try:
            file_type = file.type
            if file_type == 'application/pdf':
                content = read_pdf(file)
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                content = read_docx(file)
            elif file_type == 'text/plain':
                content = read_txt(file)

            st.header("Inhalt des Dokuments")
            input_text = st.text_area("Ausgabe", value=content, height=200, key=20)
        except:
            st.write("Add the absolute link into the text field to get a file outside of the repo")
        
    if st.button("Speech to Text", on_click=style_button_row, kwargs={
        'clicked_button_ix': 2, 'n_buttons': 2
    }):
        st.write("Click to start recording")
        
        content=record_audio(filename='transcript.txt')
       
       
        input_text = st.text_area("Speech to Text an copy the text into the Textbox", value=content, height=200, key=30)
        
        
    
    # Texteingabe
    if file is None:
        input_text = st.text_area("Enter a text", '', height=200, key=10)
        if is_file(input_text):
            print(True)
            if input_text.endswith('.docx'):
                content = read_docx_file(input_text)
                print(1)
            elif input_text.endswith('.txt'):
                content = read_txt_file(input_text)
                print(2)
                print(content)
            elif input_text.endswith('.pdf'):
                content = read_pdf_file(input_text)
                print(3)
            input_text=content
            input_text = st.text_area("Copy the text into the above text field ", input_text, height=200, key=100)
            st.write("Copy the Data into above text field ")

    

    if st.button("read text", on_click=style_button_row, kwargs={
    'clicked_button_ix': 1, 'n_buttons': 2}):
        if input_text == "":
                st.warning("Please enter a text first")
                modell_speak("Please enter a text first.")
        else:
            modell_speak(input_text)

   



    # if st.button("read text"):
    #     if input_text == "":
    #         st.warning("Please enter a text first")
    #         modell_speak("Please enter a text first.")
    #     else:
    #         modell_speak(input_text)
    # Dokument auswählen
    
    
    # Checkboxen für Klassifikation und Zusammenfassung
    classification_enabled = st.checkbox("Activate classification")
    summarization_enabled = st.checkbox("Activate summary")
    
    # Schieberegler für die Kompressionsrate
    compression_rate = st.slider("Compression rate (%) to compress the text to the value", min_value=20, max_value=80, value=50, step=1)
    
    # Knopf zum Zusammenfassen und Klassifizieren
    if st.button("Execute",key=2,type='secondary'):
        if input_text:
            try:
                input_text_neu=''
                for batch in split_into_batches(input_text):
                    input_text_neu += model_text_korrigieren.restore_punctuation(batch)
                input_text= input_text_neu
                
            except:
                pass
            if classification_enabled or summarization_enabled:
                execute_the_classification_and_summary(input_text,compression_rate,classification_enabled,summarization_enabled)
            else:
                no_box_text='Please check the Box classification and / or summary to proceed further!'
                st.warning(no_box_text)
                modell_speak(no_box_text)
        else:
            st.warning("Please enter a text first.")

    if st.button("Helper",type='primary'):
        scren_text='''Welcome to Syntex. Please input your text in the text area and press the read text button to hear the text. 
        If you want to upload a document, please select the document type and press the execute button. 
        You can also edit the text in the designated textbox.
        If you want to activate the classification and the summary, please check the corresponding boxes. 

        You can adjust the compression rate with the slider. 
        To execute the classification and the summary, please press the execute button.
        '''
        modell_speak(scren_text)
if __name__ == '__main__':
    
    main()

